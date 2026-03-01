"""
合同模板处理工具
"""
from pathlib import Path
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from decimal import Decimal


class ContractTemplateProcessor:
    """合同模板处理器"""
    
    # 支持的占位符映射
    PLACEHOLDER_MAP = {
        'supplier_name': '供应商名称',
        'supplier_code': '供应商编码',
        'supplier_contact': '联系人',
        'supplier_phone': '联系电话',
        'supplier_email': '联系邮箱',
        'supplier_address': '地址',
        'supplier_tax_id': '税号',
        'supplier_bank_name': '开户银行',
        'supplier_bank_account': '银行账号',
        'quotation_code': '报价单编号',
        'quotation_title': '报价单标题',
        'quotation_date': '报价日期',
        'quotation_valid_until': '有效期至',
        'total_amount': '总金额',
        'currency': '币种',
        'payment_terms': '付款条件',
        'delivery_terms': '交货条件',
        'delivery_days': '交货天数',
        'contract_code': '合同编号',
        'contract_title': '合同标题',
        'sign_date': '签订日期',
        'start_date': '生效日期',
        'end_date': '到期日期',
        'items_table': '报价明细表格',
    }
    
    @staticmethod
    def format_date(date_value: Optional[datetime]) -> str:
        """格式化日期"""
        if not date_value:
            return ''
        if isinstance(date_value, str):
            try:
                date_value = datetime.fromisoformat(date_value.replace('Z', '+00:00'))
            except:
                return date_value
        return date_value.strftime('%Y年%m月%d日')
    
    @staticmethod
    def format_amount(amount: Optional[Any]) -> str:
        """格式化金额"""
        if amount is None:
            return '0.00'
        if isinstance(amount, str):
            try:
                amount = Decimal(amount)
            except:
                return amount
        return f'{float(amount):,.2f}'
    
    @staticmethod
    def replace_text_in_paragraph(paragraph, old_text: str, new_text: str):
        """替换段落中的文本"""
        if old_text in paragraph.text:
            # 保存段落格式
            runs = paragraph.runs
            full_text = ''.join([run.text for run in runs])
            
            if old_text in full_text:
                # 清空段落
                paragraph.clear()
                # 按占位符分割文本
                parts = full_text.split(old_text)
                for i, part in enumerate(parts):
                    if part:
                        run = paragraph.add_run(part)
                        # 尝试保持原有格式
                        if runs:
                            run.font.name = runs[0].font.name
                            run.font.size = runs[0].font.size
                    if i < len(parts) - 1:
                        run = paragraph.add_run(new_text)
                        if runs:
                            run.font.name = runs[0].font.name
                            run.font.size = runs[0].font.size
    
    @staticmethod
    def replace_text_in_table(table, old_text: str, new_text: str):
        """替换表格中的文本"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    ContractTemplateProcessor.replace_text_in_paragraph(paragraph, old_text, new_text)
    
    @staticmethod
    def replace_items_table(doc: Document, items_data: list, placeholder: str = '{items_table}'):
        """替换报价明细表格占位符"""
        # 查找包含占位符的段落
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                # 找到占位符所在的段落，在其后插入表格
                parent = paragraph._element.getparent()
                p_index = parent.index(paragraph._element)
                
                # 创建表格
                table = doc.add_table(rows=1, cols=7)
                table.style = 'Light Grid Accent 1'
                
                # 表头
                header_cells = table.rows[0].cells
                headers = ['序号', '物料名称', '规格型号', '单位', '数量', '单价', '总价']
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    header_cells[i].paragraphs[0].runs[0].font.bold = True
                    header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 添加数据行
                for item in items_data:
                    row = table.add_row()
                    row.cells[0].text = str(item.get('sequence', ''))
                    row.cells[1].text = str(item.get('material_name', ''))
                    row.cells[2].text = str(item.get('specification', ''))
                    row.cells[3].text = str(item.get('unit', ''))
                    row.cells[4].text = str(item.get('quantity', ''))
                    row.cells[5].text = ContractTemplateProcessor.format_amount(item.get('unit_price'))
                    row.cells[6].text = ContractTemplateProcessor.format_amount(item.get('total_price'))
                
                # 移除占位符段落
                parent.remove(paragraph._element)
                break
        
        # 也在表格中查找占位符
        for table in doc.tables:
            ContractTemplateProcessor.replace_items_table_in_table(table, items_data, placeholder)
    
    @staticmethod
    def replace_items_table_in_table(table, items_data: list, placeholder: str):
        """在表格中替换报价明细表格占位符"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        # 在单元格中创建嵌套表格
                        cell.paragraphs[0].clear()
                        nested_table = cell.add_table(rows=1, cols=7)
                        
                        # 表头
                        header_cells = nested_table.rows[0].cells
                        headers = ['序号', '物料名称', '规格型号', '单位', '数量', '单价', '总价']
                        for i, header in enumerate(headers):
                            header_cells[i].text = header
                        
                        # 添加数据行
                        for item in items_data:
                            row = nested_table.add_row()
                            row.cells[0].text = str(item.get('sequence', ''))
                            row.cells[1].text = str(item.get('material_name', ''))
                            row.cells[2].text = str(item.get('specification', ''))
                            row.cells[3].text = str(item.get('unit', ''))
                            row.cells[4].text = str(item.get('quantity', ''))
                            row.cells[5].text = ContractTemplateProcessor.format_amount(item.get('unit_price'))
                            row.cells[6].text = ContractTemplateProcessor.format_amount(item.get('total_price'))
    
    @classmethod
    def generate_contract(cls, template_path: str, output_path: str, data: Dict[str, Any]) -> str:
        """
        根据模板生成合同文档
        
        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径
            data: 填充数据字典
            
        Returns:
            生成的合同文件路径
        """
        # 加载模板
        doc = Document(template_path)
        
        # 准备替换数据
        replacements = {
            '{supplier_name}': data.get('supplier_name', ''),
            '{supplier_code}': data.get('supplier_code', ''),
            '{supplier_contact}': data.get('supplier_contact', ''),
            '{supplier_phone}': data.get('supplier_phone', ''),
            '{supplier_email}': data.get('supplier_email', ''),
            '{supplier_address}': data.get('supplier_address', ''),
            '{supplier_tax_id}': data.get('supplier_tax_id', ''),
            '{supplier_bank_name}': data.get('supplier_bank_name', ''),
            '{supplier_bank_account}': data.get('supplier_bank_account', ''),
            '{quotation_code}': data.get('quotation_code', ''),
            '{quotation_title}': data.get('quotation_title', ''),
            '{quotation_date}': cls.format_date(data.get('quotation_date')),
            '{quotation_valid_until}': cls.format_date(data.get('quotation_valid_until')),
            '{total_amount}': cls.format_amount(data.get('total_amount')),
            '{currency}': data.get('currency', 'CNY'),
            '{payment_terms}': data.get('payment_terms', ''),
            '{delivery_terms}': data.get('delivery_terms', ''),
            '{delivery_days}': str(data.get('delivery_days', '')),
            '{contract_code}': data.get('contract_code', ''),
            '{contract_title}': data.get('contract_title', ''),
            '{sign_date}': cls.format_date(data.get('sign_date')),
            '{start_date}': cls.format_date(data.get('start_date')),
            '{end_date}': cls.format_date(data.get('end_date')),
        }
        
        # 替换段落中的文本
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                cls.replace_text_in_paragraph(paragraph, old_text, new_text)
        
        # 替换表格中的文本
        for table in doc.tables:
            for old_text, new_text in replacements.items():
                cls.replace_text_in_table(table, old_text, new_text)
        
        # 处理报价明细表格
        items_data = data.get('items', [])
        if items_data:
            cls.replace_items_table(doc, items_data)
        
        # 保存文档
        doc.save(output_path)
        return output_path

